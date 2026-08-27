import logging
import os
import re
from datetime import datetime, timedelta
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, ContextTypes, CommandHandler, MessageHandler, CallbackQueryHandler, filters
from docx import Document
import openpyxl
from flask import Flask, request
import asyncio

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)

app_flask = Flask(__name__)

TOKEN = "8988591586:AAFdWPkI7MGaJcAmoclzbAn9lkKXgarS6z4"
WEBHOOK_SECRET = "my_secret_token_123"

telegram_app = ApplicationBuilder().token(TOKEN).build()

@app_flask.route('/')
def home():
    return "Telegram Bot is running smoothly with Webhook!"

@app_flask.route(f'/{TOKEN}', methods=['POST'])
def webhook():
    if request.headers.get('X-Telegram-Bot-Api-Secret-Token') == WEBHOOK_SECRET:
        update = Update.de_json(request.get_json(force=True), telegram_app.bot)
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        loop.run_until_complete(telegram_app.process_update(update))
        loop.close()
        return 'OK', 200
    return 'Unauthorized', 403

user_last_files = {}
user_meetings = {}        

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROCESSED_FOLDER = os.path.join(BASE_DIR, "processed_files")

if not os.path.exists(PROCESSED_FOLDER):
    os.makedirs(PROCESSED_FOLDER)

KHMER_DICTIONARY_CORRECTIONS = {
    "សំលាប់": "សម្លាប់",
    "សំលៀកបំពាក់": "សម្លៀកបំពាក់",
    "សំលេង": "សម្លេង",
    "សំបុត្រ": "សំបុត្រ",
    "ត្រួតពិនត្យ": "ត្រួតពិនិត្យ",
    "អនុញ្ញាត្តិ": "អនុញ្ញាត",
    "កំរិត": "កម្រិត",
    "កំលាំង": "កម្លាំង",
    "ទំនេរ": "ទំនេរ",
    "ដំនើរការ": "ដំណើរការ",
    "ចំនុច": "ចំណុច",
    "ចំនាយ": "ចំណាយ",
    "ចំនូល": "ចំណូល",
    "រៀបចំរៀង": "រៀបរៀង",
    "ពត៌មាន": "ព័ត៌មាន",
    "ឧបករណ៏": "ឧបករណ៍"
}

def clean_khmer_text(text):
    if not text:
        return text
    for wrong, correct in KHMER_DICTIONARY_CORRECTIONS.items():
        text = text.replace(wrong, correct)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\s+([,.\?!៖៕៕])', r'\1', text)
    return text.strip()

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_name = update.effective_user.first_name
    welcome_message = (
        f"សួស្តី {user_name}! 🙏\n"
        "ខ្ញុំជា Bot ជំនួយការការងារ (My Work Assistant)។\n\n"
        "📁 **Attach ហ្វាល Word** - សម្រាប់គ្រប់គ្រងហ្វាលតាមប៊ូតុងអន្តរកម្ម។\n"
        "📅 `/meeting` - កត់ត្រាកាលវិភាគប្រជុំតាមទម្រង់ខាងក្រោម៖\n\n"
        "📌 **ឧទាហរណ៍៖**\n"
        "`/meeting ប្រជុំយុទ្ធសាស្ត្រលក់ដូរ | 28-08-2026 14:30 | 15`"
    )
    await update.message.reply_text(welcome_message, parse_mode="Markdown")

async def meeting_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text_args = " ".join(context.args)

    if not text_args or "|" not in text_args:
        await update.message.reply_text(
            "⚠️ **ទម្រង់នៃការប្រើប្រាស់មិនទាន់ត្រឹមត្រូវទេ!**\n\n"
            "សូមប្រើប្រាស់តាមទម្រង់នេះ៖\n"
            "`/meeting ប្រធានបទ | DD-MM-YYYY HH:MM | នាទីរំលឹក`\n\n"
            "📌 **ឧទាហរណ៍៖**\n"
            "`/meeting ប្រជុំយុទ្ធសាស្ត្រលក់ដូរ | 28-08-2026 14:30 | 15`",
            parse_mode="Markdown"
        )
        return

    try:
        parts = [p.strip() for p in text_args.split("|")]
        if len(parts) < 3:
            await update.message.reply_text("⚠️ សូមបំពេញព័ត៌មានឱ្យបានគ្រប់គ្រង ៣ ចំណែក (ប្រធានបទ | ថ្ងៃម៉ោង | នាទីរំលឹក)")
            return

        topic = parts[0]
        datetime_str = parts[1]
        remind_mins = int(parts[2])

        target_dt = datetime.strptime(datetime_str, "%d-%m-%Y %H:%M")
        now = datetime.now()
        if target_dt <= now:
            await update.message.reply_text("❌ ថ្ងៃ និងម៉ោងបានន្លងផុតទៅហើយ។ សូមកំណត់ពេលថ្ងៃមុខ។")
            return

        if user_id not in user_meetings:
            user_meetings[user_id] = []

        new_meeting = {
            "topic": topic,
            "date_time": datetime_str,
            "remind": f"{remind_mins} នាទីមុន"
        }
        user_meetings[user_id].append(new_meeting)

        date_part = datetime_str.split(' ')[0]
        time_part = datetime_str.split(' ')[1]
        
        report_text = (
            f"🎉 **កត់ត្រាកាលវិភាគប្រជុំជោគជ័យ!**\n\n"
            f"គោរពរាយការណ៍ជូនមេ! ថ្ងៃនេះមានការប្រជុំ{topic} ថ្ងៃ {date_part} វេលាម៉ោង {time_part} (រំលឹកមុន {remind_mins} នាទី)"
        )
        await update.message.reply_text(report_text, parse_mode="Markdown")

    except Exception as e:
        await update.message.reply_text("❌ មានបញ្ហា៖ សូមពិនិត្យទម្រង់ថ្ងៃខែឆ្នាំ និងតួលេខនាទីឡើងវិញ។")

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    data = query.data

    if data.startswith("btn_"):
        if user_id not in user_last_files or not os.path.exists(user_last_files[user_id]):
            await query.message.reply_text("⚠️ សូម Attach ហ្វាល Word មកជាមុនសិន។")
            return
        target_file = user_last_files[user_id]
        
        try:
            if data == "btn_summary":
                doc = Document(target_file)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                summary_items = paragraphs[:4]
                summary_text = "\n".join([f"• {item}" for item in summary_items])
                if not summary_text:
                    summary_text = "(ហ្វាលនេះទទេ ឬគ្មានអត្ថបទច្បាស់លាស់)"
                file_name = os.path.basename(target_file)
                await query.message.reply_text(
                    f"📋 **សេចក្តីសង្ខេបរបាយការណ៍ (`{file_name}`)៖**\n"
                    f"-----------------------------------\n"
                    f"{summary_text}\n"
                    f"-----------------------------------",
                    parse_mode="Markdown"
                )
            elif data in ["btn_preview1", "btn_preview2"]:
                page_num = 1 if data == "btn_preview1" else 2
                doc = Document(target_file)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                page_size = 12
                total_pages = (len(paragraphs) + page_size - 1) // page_size
                if page_num > total_pages:
                    await query.message.reply_text(f"⚠️ ឯកសារនេះមានត្រឹមតែ **ទំព័រទី {total_pages}** ប៉ុណ្ណោះ។")
                    return
                start_idx = (page_num - 1) * page_size
                end_idx = start_idx + page_size
                page_content = "\n\n".join(paragraphs[start_idx:end_idx])
                await query.message.reply_text(
                    f"📄 **[ទំព័រទី {page_num} នៃ {total_pages}]**\n"
                    f"-----------------------------------\n"
                    f"{page_content}\n"
                    f"-----------------------------------",
                    parse_mode="Markdown"
                )
            elif data == "btn_spellcheck":
                doc = Document(target_file)
                fixed_count = 0
                for p in doc.paragraphs:
                    for run in p.runs:
                        orig = run.text
                        cleaned = clean_khmer_text(orig)
                        if orig != cleaned:
                            run.text = cleaned
                            fixed_count += 1
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    orig = run.text
                                    cleaned = clean_khmer_text(orig)
                                    if orig != cleaned:
                                        run.text = cleaned
                                        fixed_count += 1
                original_base_name = os.path.basename(target_file)
                fixed_file_name = os.path.join(PROCESSED_FOLDER, original_base_name)
                doc.save(fixed_file_name)
                await query.message.reply_text(f"✨ បានកែអក្ខរាវិរុទ្ធ និងដកឃ្លាជោគជ័យចំនួន `{fixed_count}` កន្លែង។", parse_mode="Markdown")
                await query.message.reply_document(document=open(fixed_file_name, "rb"), caption="✅ ហ្វាល Word ស្អាត រក្សា Font ដើម!")
            elif data == "btn_editword_info":
                await query.message.reply_text("✏️ សូមប្រើពាក្យបញ្ជា៖ `/editword ពាក្យចាស់ | ពាក្យថ្មី`")
        except Exception as e:
            await query.message.reply_text(f"❌ មានបញ្ហា៖ {str(e)}")

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        user_id = update.effective_user.id
        document = update.message.document
        file_name = document.file_name
        
        saved_path = os.path.join(BASE_DIR, file_name)
        file = await context.bot.get_file(document.file_id)
        await file.download_to_drive(saved_path)
        user_last_files[user_id] = saved_path
        
        if file_name.endswith('.docx'):
            doc = Document(saved_path)
            total_paragraphs = len([p.text for p in doc.paragraphs if p.text.strip()])
            
            keyboard = [
                [InlineKeyboardButton("📋 សង្ខេប (Summary)", callback_data="btn_summary")],
                [InlineKeyboardButton("📄 អានទំព័រទី១ (Preview 1)", callback_data="btn_preview1"),
                 InlineKeyboardButton("📄 អានទំព័រទី២ (Preview 2)", callback_data="btn_preview2")],
                [InlineKeyboardButton("✍️ កែអក្ខរាវិរុទ្ធ (Spellcheck)", callback_data="btn_spellcheck")],
                [InlineKeyboardButton("✏️ កែពាក្យ (Editword)", callback_data="btn_editword_info")]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            await update.message.reply_text(
                f"📥 **ទទួលបាន និងចងចាំឯកសារ Word (`{file_name}`) រួចរាល់!**\n"
                f"📊 ចំនួនកថាខណ្ឌសរុប៖ {total_paragraphs} ឃ្លា\n\n"
                f"👇 **សូមចុចជ្រើសរើសមុខងារខាងក្រោមនេះ៖**",
                reply_markup=reply_markup,
                parse_mode="Markdown"
            )
        else:
            await update.message.reply_text(f"📥 ទទួលបានឯកសារ `{file_name}` រួចរាល់!")
    except Exception as e:
        await update.message.reply_text(f"❌ មានបញ្ហា៖ {str(e)}")

def setup_handlers():
    # ចុះឈ្មោះ Command Handler ទាំងពីរឱ្យច្បាស់លាស់
    telegram_app.add_handler(CommandHandler("start", start))
    telegram_app.add_handler(CommandHandler("meeting", meeting_command))
    telegram_app.add_handler(CallbackQueryHandler(button_handler))
    telegram_app.add_handler(MessageHandler(filters.Document.ALL, handle_document))

if __name__ == '__main__':
    setup_handlers()

    # ធ្វើការ Initialize និង Start Telegram App ទុកជាមុនសិន ធានាថា Handlers ទាំងអស់ដំណើរការ
    asyncio.run(telegram_app.initialize())
    asyncio.run(telegram_app.start())

    # បើកដំណើរការ Flask Server ទទួល Webhook
    port = int(os.environ.get("PORT", 10000))
    app_flask.run(host="0.0.0.0", port=port)
